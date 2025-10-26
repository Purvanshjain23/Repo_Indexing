#!/usr/bin/env python3
"""
Batch HTML to DOCX Converter
Converts all HTML files from HTML_Outputs to Claude_Docx folder
Uses enhanced professional converter with emojis and optimization
"""

import os
import sys
import base64
import io
import tempfile
import re
from collections import defaultdict
from pathlib import Path
from bs4 import BeautifulSoup, NavigableString
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import requests
from PIL import Image
import time

class EnhancedConverter:
    """Enhanced Professional Converter - same as enhanced_professional_converter.py"""

    def __init__(self):
        self.doc = Document()
        self.mermaid_counter = 0
        self.table_counter = 0
        self.processed_elements = set()
        self.setup_professional_styles()

    def setup_professional_styles(self):
        """Setup professional document styles"""
        styles = self.doc.styles
        normal = styles['Normal']
        normal.font.name = 'Calibri'
        normal.font.size = Pt(11)

        for i in range(1, 7):
            try:
                heading = styles[f'Heading {i}']
                heading.font.color.rgb = RGBColor(0, 102, 204)
            except:
                pass

    def add_emoji_to_heading(self, text):
        """Add contextual emojis to headings"""
        emoji_map = {
            'executive summary': '📊', 'business': '🎯', 'value': '💎',
            'system scale': '📈', 'operational impact': '⚡', 'performance': '⚡',
            'table of contents': '📑', 'program context': '🎯', 'inputs': '📥',
            'structure': '🏗️', 'logic': '🔄', 'data interaction': '💾',
            'dependencies': '🔗', 'modernization': '🚀', 'functions': '⚙️',
            'subroutine': '🔧', 'analysis': '📊', 'overview': '🔍',
            'documentation': '📖', 'inventory': '📋', 'retrieval': '🔍',
            'processing': '⚙️', 'validation': '✓', 'integration': '🔗',
        }

        text_lower = text.lower()
        for keyword, emoji in emoji_map.items():
            if keyword in text_lower:
                return f"{emoji} {text}"
        return text

    def render_mermaid(self, code, num):
        """Render mermaid diagram"""
        try:
            print(f"    [{num}] Rendering...", end='', flush=True)
            encoded = base64.urlsafe_b64encode(code.strip().encode('utf-8')).decode('utf-8')
            url = f"https://mermaid.ink/img/{encoded}"

            for attempt in range(2):
                try:
                    response = requests.get(url, timeout=30)
                    if response.status_code == 200:
                        img = Image.open(io.BytesIO(response.content))
                        print(" OK")
                        return img
                    time.sleep(1)
                except:
                    if attempt < 1:
                        time.sleep(2)

            print(" FAILED")
        except:
            print(" ERROR")
        return None

    def optimize_image(self, img):
        """Optimize image"""
        if img.mode in ('RGBA', 'LA', 'P'):
            bg = Image.new('RGB', img.size, (255, 255, 255))
            if img.mode == 'RGBA':
                bg.paste(img, mask=img.split()[-1])
            else:
                bg.paste(img)
            img = bg

        if img.width > 1000:
            ratio = 1000 / img.width
            img = img.resize((int(img.width * ratio), int(img.height * ratio)), Image.Resampling.LANCZOS)

        return img

    def add_optimized_image(self, img, caption=None):
        """Add optimized image"""
        try:
            with tempfile.NamedTemporaryFile(delete=False, suffix='.jpg') as tmp:
                tmp_name = tmp.name
                img.save(tmp_name, 'JPEG', quality=70, optimize=True)

            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(tmp_name, width=Inches(5.5))

            if caption:
                cap_p = self.doc.add_paragraph(caption)
                cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in cap_p.runs:
                    run.font.size = Pt(9)
                    run.font.italic = True
                    run.font.color.rgb = RGBColor(128, 128, 128)

            try:
                os.unlink(tmp_name)
            except:
                pass
        except Exception as e:
            print(f"    Image error: {str(e)[:30]}")

    def add_styled_table(self, table_element):
        """Add table"""
        try:
            self.table_counter += 1
            rows = table_element.find_all('tr')
            if not rows:
                return

            max_cols = max(len(row.find_all(['th', 'td'])) for row in rows)
            doc_table = self.doc.add_table(rows=len(rows), cols=max_cols)
            doc_table.style = 'Light Grid Accent 1'

            for i, row in enumerate(rows):
                cells = row.find_all(['th', 'td'])
                for j, cell in enumerate(cells):
                    if j < max_cols:
                        doc_table.rows[i].cells[j].text = cell.get_text(strip=True)

                        if cell.name == 'th' or i == 0:
                            for para in doc_table.rows[i].cells[j].paragraphs:
                                for run in para.runs:
                                    run.font.bold = True
                                    run.font.color.rgb = RGBColor(255, 255, 255)

                            cell_elem = doc_table.rows[i].cells[j]._element
                            tcPr = cell_elem.get_or_add_tcPr()
                            shading = OxmlElement('w:shd')
                            shading.set(qn('w:fill'), '0066cc')
                            tcPr.append(shading)

            self.doc.add_paragraph()
        except Exception as e:
            print(f"    Table error: {str(e)[:30]}")

    def extract_function_names(self, soup):
        """Extract function names"""
        text = soup.get_text()
        patterns = [
            r'\b([A-Z]{2,}[A-Z0-9]{2,})\s*[-:]',
            r'\b([A-Z]{2,}[A-Z0-9]{2,})\s*\(',
            r'\b([A-Z]{2,}RVGN|[A-Z]{2,}CHRC|[A-Z]{2,}SUBR)\b',
        ]

        functions = set()
        for pattern in patterns:
            matches = re.findall(pattern, text)
            functions.update(matches)

        return sorted(list(functions))

    def create_function_inventory_table(self, functions, title):
        """Create function inventory table"""
        if not functions:
            return

        h = self.doc.add_heading(f"📋 {title}", 3)
        h.runs[0].font.color.rgb = RGBColor(0, 102, 204)

        num_funcs = min(len(functions), 50)  # Limit to 50 per table
        doc_table = self.doc.add_table(rows=num_funcs + 1, cols=4)
        doc_table.style = 'Light Grid Accent 1'

        headers = ['Function', 'Purpose', 'Type', 'Category']
        for j, header in enumerate(headers):
            cell = doc_table.rows[0].cells[j]
            cell.text = header
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)

            cell_elem = cell._element
            tcPr = cell_elem.get_or_add_tcPr()
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), '0066cc')
            tcPr.append(shading)

        for i, func in enumerate(functions[:num_funcs], 1):
            doc_table.rows[i].cells[0].text = func
            doc_table.rows[i].cells[1].text = self.infer_purpose(func)
            doc_table.rows[i].cells[2].text = self.infer_type(func)
            doc_table.rows[i].cells[3].text = self.infer_category(func)

        self.doc.add_paragraph()

    def infer_purpose(self, func):
        """Infer function purpose"""
        if 'RVGN' in func:
            return 'Data retrieval and validation'
        elif 'CHRC' in func:
            return 'Change record processing'
        elif 'SUBR' in func:
            return 'Utility subroutine'
        elif func.startswith('ZZ'):
            return 'System initialization'
        elif func.startswith('BA') or func.startswith('BB'):
            return 'Subfile management'
        elif func.startswith('EB') or func.startswith('DC'):
            return 'Primary business logic'
        else:
            return 'Business processing'

    def infer_type(self, func):
        """Infer function type"""
        if 'RVGN' in func:
            return 'Retrieval'
        elif 'CHRC' in func:
            return 'Update'
        elif 'SUBR' in func:
            return 'Utility'
        else:
            return 'Processing'

    def infer_category(self, func):
        """Infer category"""
        if 'RVGN' in func:
            return 'Data Access'
        elif 'CHRC' in func:
            return 'Data Change'
        elif func.startswith('ZZ'):
            return 'Initialization'
        elif func.startswith('BA') or func.startswith('BB'):
            return 'UI Management'
        else:
            return 'Business Logic'

    def categorize_functions(self, functions):
        """Categorize functions"""
        categories = {
            'Data Retrieval': [],
            'Change Processing': [],
            'Initialization': [],
            'Subfile Management': [],
            'Business Logic': [],
            'Utilities': [],
        }

        for func in functions:
            if 'RVGN' in func:
                categories['Data Retrieval'].append(func)
            elif 'CHRC' in func:
                categories['Change Processing'].append(func)
            elif func.startswith('ZZ'):
                categories['Initialization'].append(func)
            elif func.startswith('BA') or func.startswith('BB'):
                categories['Subfile Management'].append(func)
            elif 'SUBR' in func:
                categories['Utilities'].append(func)
            else:
                categories['Business Logic'].append(func)

        return categories

    def process_content_enhanced(self, soup):
        """Process content maintaining structure"""
        sections = soup.find_all('section')

        for section in sections:
            for element in section.descendants:
                if not hasattr(element, 'name') or not element.name:
                    continue

                elem_id = id(element)
                if elem_id in self.processed_elements:
                    continue

                elem_name = element.name

                if elem_name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                    text = element.get_text(strip=True)
                    if text and len(text) > 2:
                        level = int(elem_name[1])
                        text_with_emoji = self.add_emoji_to_heading(text)
                        h = self.doc.add_heading(text_with_emoji, level=level)
                        h.runs[0].font.color.rgb = RGBColor(0, 102, 204)
                        self.processed_elements.add(elem_id)

                elif elem_name == 'p':
                    text = element.get_text(strip=True)
                    if text and len(text) > 5 and element.parent.name not in ['li', 'td', 'th']:
                        self.doc.add_paragraph(text)
                        self.processed_elements.add(elem_id)

                elif elem_name == 'ul':
                    items = element.find_all('li', recursive=False)
                    if items and elem_id not in self.processed_elements:
                        for item in items:
                            text = item.get_text(strip=True)
                            if text:
                                self.doc.add_paragraph(text, style='List Bullet')
                        self.processed_elements.add(elem_id)

                elif elem_name == 'ol':
                    items = element.find_all('li', recursive=False)
                    if items and elem_id not in self.processed_elements:
                        for item in items:
                            text = item.get_text(strip=True)
                            if text:
                                self.doc.add_paragraph(text, style='List Number')
                        self.processed_elements.add(elem_id)

                elif elem_name == 'div' and 'mermaid' in element.get('class', []):
                    if elem_id not in self.processed_elements:
                        self.mermaid_counter += 1
                        code = element.get_text(strip=True)

                        p = self.doc.add_paragraph()
                        run = p.add_run(f"📊 Figure {self.mermaid_counter}: Process Flow Diagram")
                        run.bold = True
                        run.font.color.rgb = RGBColor(0, 102, 204)
                        run.font.size = Pt(12)

                        img = self.render_mermaid(code, self.mermaid_counter)
                        if img:
                            img = self.optimize_image(img)
                            self.add_optimized_image(img, f"Diagram {self.mermaid_counter}")
                        else:
                            fallback = self.doc.add_paragraph("[!] Diagram rendering failed")
                            fallback.runs[0].font.italic = True
                            fallback.runs[0].font.color.rgb = RGBColor(192, 57, 43)

                        self.processed_elements.add(elem_id)

                elif elem_name == 'table':
                    if elem_id not in self.processed_elements:
                        self.add_styled_table(element)
                        self.processed_elements.add(elem_id)

    def convert_file(self, html_file, output_file):
        """Convert single HTML file"""
        # Reset for new file
        self.doc = Document()
        self.mermaid_counter = 0
        self.table_counter = 0
        self.processed_elements = set()
        self.setup_professional_styles()

        # Read
        with open(html_file, 'r', encoding='utf-8') as f:
            html_content = f.read()

        # Parse
        soup = BeautifulSoup(html_content, 'html.parser')

        # Title
        title = soup.find('title')
        if title:
            self.doc.add_heading(f"📋 {title.get_text(strip=True)}", 0)

        # Extract functions
        functions = self.extract_function_names(soup)

        # Find elements
        sections = soup.find_all('section')
        mermaid_divs = soup.find_all('div', class_='mermaid')
        tables = soup.find_all('table')

        # Process content
        self.process_content_enhanced(soup)

        # Add function inventory
        if functions and len(functions) > 5:
            categories = self.categorize_functions(functions)
            for category, funcs in categories.items():
                if funcs:
                    self.create_function_inventory_table(funcs, f"{category} Functions")

        # Save
        self.doc.save(output_file)

        return {
            'diagrams': f"{self.mermaid_counter}/{len(mermaid_divs)}",
            'tables': self.table_counter,
            'functions': len(functions),
            'size': os.path.getsize(output_file)
        }


def batch_convert_html_to_docx(input_dir, output_dir):
    """Convert all HTML files in input directory"""
    print("\n" + "=" * 70)
    print("BATCH HTML TO DOCX CONVERTER")
    print("=" * 70)

    input_path = Path(input_dir)
    output_path = Path(output_dir)

    # Find all HTML files
    html_files = sorted(input_path.glob('*.html'))

    if not html_files:
        print(f"\nNo HTML files found in {input_dir}")
        return

    print(f"\nFound {len(html_files)} HTML files to convert:")
    for i, html_file in enumerate(html_files, 1):
        print(f"  {i}. {html_file.name}")

    print("\n" + "=" * 70)
    print("STARTING CONVERSIONS")
    print("=" * 70 + "\n")

    converter = EnhancedConverter()
    results = []

    for i, html_file in enumerate(html_files, 1):
        print(f"[{i}/{len(html_files)}] Converting {html_file.name}...")

        output_file = output_path / f"{html_file.stem}.docx"

        try:
            stats = converter.convert_file(str(html_file), str(output_file))
            results.append({
                'file': html_file.name,
                'output': output_file.name,
                'status': 'SUCCESS',
                **stats
            })
            print(f"  SUCCESS: Saved {output_file.name}")
            print(f"    Diagrams: {stats['diagrams']}, Tables: {stats['tables']}, Size: {stats['size']:,} bytes")

        except Exception as e:
            print(f"  ERROR: {str(e)[:60]}")
            results.append({
                'file': html_file.name,
                'status': 'FAILED',
                'error': str(e)[:60]
            })

        print()

    # Summary
    print("=" * 70)
    print("BATCH CONVERSION SUMMARY")
    print("=" * 70)

    successful = sum(1 for r in results if r['status'] == 'SUCCESS')
    failed = sum(1 for r in results if r['status'] == 'FAILED')

    print(f"\nTotal files: {len(html_files)}")
    print(f"Successful: {successful}")
    print(f"Failed: {failed}")

    if successful > 0:
        print(f"\nConverted files:")
        for r in results:
            if r['status'] == 'SUCCESS':
                print(f"  [OK] {r['file']:40} -> {r['output']}")
                print(f"       Stats: {r.get('diagrams', 'N/A')} diagrams, {r.get('tables', 0)} tables, {r.get('size', 0)/1024:.1f} KB")

    if failed > 0:
        print(f"\nFailed files:")
        for r in results:
            if r['status'] == 'FAILED':
                print(f"  [FAIL] {r['file']:40} - {r.get('error', 'Unknown error')}")

    print("\n" + "=" * 70 + "\n")


if __name__ == "__main__":
    input_dir = r"C:\Users\PurvanshJain\OneDrive - Programmers.IO\Desktop\Full_Repo_Analysis\HTML_Outputs"
    output_dir = r"C:\Users\PurvanshJain\OneDrive - Programmers.IO\Desktop\Full_Repo_Analysis\Claude_Docx"

    batch_convert_html_to_docx(input_dir, output_dir)
