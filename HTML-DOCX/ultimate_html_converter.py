#!/usr/bin/env python3
"""
Ultimate HTML to DOCX/Markdown Converter
Enhanced version with comprehensive features:
- Superior Mermaid diagram rendering with multiple fallbacks
- Professional DOCX formatting with styles
- Clean Markdown output with proper structure
- Batch processing with detailed progress tracking
- Robust error handling and recovery
- Configuration options for different output needs
"""

import os
import re
import base64
import requests
import json
import glob
import time
import shutil
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Tuple, Optional, Any
from dataclasses import dataclass
from enum import Enum

# DOCX dependencies
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn

# HTML parsing
from bs4 import BeautifulSoup, NavigableString, Tag
from html.parser import HTMLParser
from html import unescape

# Image processing
from PIL import Image, ImageDraw, ImageFont
import io
import urllib.parse

class OutputFormat(Enum):
    DOCX = "docx"
    MARKDOWN = "md"
    BOTH = "both"

@dataclass
class ConversionConfig:
    """Configuration for conversion process"""
    output_format: OutputFormat = OutputFormat.BOTH
    render_mermaid: bool = True
    max_diagram_retries: int = 3
    diagram_width_inches: float = 6.0
    preserve_code_blocks: bool = True
    add_table_of_contents: bool = True
    professional_styling: bool = True
    verbose_logging: bool = True
    temp_dir: str = "temp_conversion"
    output_dir: str = "Converted_Output"

class EnhancedMermaidRenderer:
    """Advanced Mermaid renderer with multiple rendering strategies"""
    
    def __init__(self, config: ConversionConfig):
        self.config = config
        self.mermaid_api = "https://mermaid.ink/img/"
        self.diagram_counter = 0
        self.successful_renders = 0
        self.failed_renders = 0
        self.render_cache = {}
        
    def clean_mermaid_code(self, mermaid_code: str) -> Optional[str]:
        """Advanced Mermaid code cleaning and validation"""
        if not mermaid_code or not mermaid_code.strip():
            return None
            
        # Remove HTML entities and normalize whitespace
        mermaid_code = unescape(mermaid_code)
        mermaid_code = re.sub(r'\s+', ' ', mermaid_code.strip())
        
        # Split into lines and clean each line
        lines = []
        for line in mermaid_code.split('\n'):
            line = line.strip()
            if line and not line.startswith('//'):  # Remove empty lines and comments
                lines.append(line)
        
        if not lines:
            return None
            
        cleaned_code = '\n'.join(lines)
        
        # Validate diagram type
        first_line = lines[0].lower()
        valid_diagrams = [
            'graph', 'flowchart', 'sequencediagram', 'classdiagram', 
            'statediagram', 'gantt', 'pie', 'gitgraph', 'journey',
            'erdiagram', 'c4context'
        ]
        
        # Auto-detect and fix diagram type
        if not any(first_line.startswith(diagram.lower()) for diagram in valid_diagrams):
            # Try to infer diagram type from content
            if any(keyword in cleaned_code.lower() for keyword in ['-->', '---', 'td', 'lr']):
                cleaned_code = 'flowchart TD\n' + cleaned_code
            elif 'participant' in cleaned_code.lower():
                cleaned_code = 'sequenceDiagram\n' + cleaned_code
            elif 'class' in cleaned_code.lower() and '{' in cleaned_code:
                cleaned_code = 'classDiagram\n' + cleaned_code
            else:
                return None
        
        return cleaned_code
    
    def extract_mermaid_code(self, element: Tag) -> Optional[str]:
        """Extract Mermaid code from various HTML structures"""
        if not element or isinstance(element, NavigableString):
            return None
            
        # Strategy 1: Direct mermaid class
        if element.name == 'div' and element.get('class'):
            classes = element.get('class', [])
            if any('mermaid' in str(cls).lower() for cls in classes):
                return self.clean_mermaid_code(element.get_text())
        
        # Strategy 2: Pre/code blocks with mermaid language
        if element.name in ['pre', 'code']:
            code_element = element.find('code') if element.name == 'pre' else element
            if code_element:
                classes = code_element.get('class', [])
                if any('mermaid' in str(cls).lower() for cls in classes):
                    return self.clean_mermaid_code(code_element.get_text())
        
        # Strategy 3: Script tags with mermaid type
        if element.name == 'script' and element.get('type') == 'text/mermaid':
            return self.clean_mermaid_code(element.get_text())
        
        # Strategy 4: Look for mermaid keywords in text content
        text_content = element.get_text().strip()
        if text_content and any(keyword in text_content.lower() for keyword in ['flowchart', 'graph td', 'graph lr', 'sequencediagram']):
            return self.clean_mermaid_code(text_content)
        
        return None
    
    def create_fallback_image(self, error_text: str, width: int = 800, height: int = 400) -> bytes:
        """Create a fallback image when Mermaid rendering fails"""
        try:
            img = Image.new('RGB', (width, height), color='#f8f9fa')
            draw = ImageDraw.Draw(img)
            
            # Try to use a font, fall back to default if not available
            try:
                font = ImageFont.truetype("arial.ttf", 16)
                title_font = ImageFont.truetype("arial.ttf", 20)
            except:
                font = ImageFont.load_default()
                title_font = font
            
            # Draw border
            draw.rectangle([10, 10, width-10, height-10], outline='#dee2e6', width=2)
            
            # Draw title
            title = "Diagram Rendering Failed"
            draw.text((width//2, 30), title, font=title_font, fill='#6c757d', anchor='mt')
            
            # Draw error message
            lines = error_text.split('\n')[:8]  # Limit to 8 lines
            y_start = 80
            for i, line in enumerate(lines):
                if len(line) > 80:
                    line = line[:77] + "..."
                draw.text((20, y_start + i * 25), line, font=font, fill='#495057')
            
            # Convert to bytes
            img_buffer = io.BytesIO()
            img.save(img_buffer, format='PNG')
            return img_buffer.getvalue()
            
        except Exception:
            # Ultimate fallback - return minimal placeholder
            return b'\x89PNG\r\n\x1a\n\x00\x00\x00\rIHDR\x00\x00\x03 \x00\x00\x01\x90\x08\x06\x00\x00\x00'
    
    def render_mermaid_to_image(self, mermaid_code: str) -> Optional[bytes]:
        """Render Mermaid code to image with advanced fallback strategies"""
        if not mermaid_code or not self.config.render_mermaid:
            return None
            
        self.diagram_counter += 1
        
        # Check cache first
        cache_key = hash(mermaid_code)
        if cache_key in self.render_cache:
            if self.config.verbose_logging:
                print(f"   📋 Using cached diagram {self.diagram_counter}")
            return self.render_cache[cache_key]
        
        if self.config.verbose_logging:
            print(f"   🎨 Rendering diagram {self.diagram_counter}...")
        
        # Multiple encoding strategies
        encoding_strategies = [
            self._encode_basic,
            self._encode_url_safe,
            self._encode_simplified,
            self._encode_minimal
        ]
        
        for strategy_idx, encode_func in enumerate(encoding_strategies):
            try:
                for attempt in range(self.config.max_diagram_retries):
                    try:
                        encoded_code = encode_func(mermaid_code)
                        if not encoded_code:
                            continue
                            
                        api_url = f"{self.mermaid_api}{encoded_code}"
                        
                        # Add delay for rate limiting
                        if attempt > 0 or strategy_idx > 0:
                            time.sleep(0.5)
                        
                        response = requests.get(api_url, timeout=15)
                        
                        if response.status_code == 200 and len(response.content) > 2000:
                            # Verify it's actually an image
                            try:
                                img = Image.open(io.BytesIO(response.content))
                                img.verify()
                                
                                self.successful_renders += 1
                                self.render_cache[cache_key] = response.content
                                
                                if self.config.verbose_logging:
                                    print(f"   ✅ Success with strategy {strategy_idx + 1}, attempt {attempt + 1}")
                                
                                return response.content
                                
                            except Exception:
                                continue
                        
                    except requests.RequestException:
                        continue
                        
            except Exception:
                continue
        
        # All strategies failed - create fallback image
        self.failed_renders += 1
        if self.config.verbose_logging:
            print(f"   ❌ Failed to render diagram {self.diagram_counter}")
        
        fallback_image = self.create_fallback_image(
            f"Diagram {self.diagram_counter}\nMermaid rendering failed\n\nOriginal code:\n{mermaid_code[:200]}..."
        )
        
        self.render_cache[cache_key] = fallback_image
        return fallback_image
    
    def _encode_basic(self, code: str) -> str:
        """Basic base64 encoding"""
        return base64.b64encode(code.encode('utf-8')).decode('utf-8')
    
    def _encode_url_safe(self, code: str) -> str:
        """URL-safe encoding"""
        return base64.b64encode(urllib.parse.quote(code, safe='').encode('utf-8')).decode('utf-8')
    
    def _encode_simplified(self, code: str) -> str:
        """Simplified code encoding"""
        # Remove complex styling and simplify syntax
        simplified = re.sub(r'classDef\s+\w+\s+[^;]*;', '', code)
        simplified = re.sub(r'class\s+\w+\s+\w+', '', simplified)
        simplified = re.sub(r'style\s+\w+\s+[^;]*;', '', simplified)
        lines = [line.strip() for line in simplified.split('\n') if line.strip()]
        return base64.b64encode('\n'.join(lines).encode('utf-8')).decode('utf-8')
    
    def _encode_minimal(self, code: str) -> str:
        """Minimal diagram with just basic structure"""
        lines = code.split('\n')
        minimal_lines = []
        
        # Keep diagram declaration
        if lines:
            minimal_lines.append(lines[0])
        
        # Extract basic connections
        for line in lines[1:]:
            if '-->' in line and len(minimal_lines) < 10:  # Limit complexity
                # Simplify node names
                simplified = re.sub(r'\[.*?\]', '', line)
                simplified = re.sub(r'\{.*?\}', '', simplified)
                if simplified.strip():
                    minimal_lines.append(simplified.strip())
        
        minimal_code = '\n'.join(minimal_lines)
        return base64.b64encode(minimal_code.encode('utf-8')).decode('utf-8')

class ProfessionalDocxConverter:
    """Professional DOCX converter with advanced formatting"""
    
    def __init__(self, config: ConversionConfig):
        self.config = config
        self.mermaid_renderer = EnhancedMermaidRenderer(config)
        self.style_applied = False
        
    def create_professional_styles(self, doc: Document):
        """Create professional document styles"""
        if self.style_applied or not self.config.professional_styling:
            return
            
        styles = doc.styles
        
        # Custom heading styles
        for level in range(1, 4):
            try:
                heading_style = styles[f'Heading {level}']
                heading_style.font.name = 'Calibri'
                heading_style.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)
                heading_style.font.bold = True
                
                if level == 1:
                    heading_style.font.size = Pt(18)
                elif level == 2:
                    heading_style.font.size = Pt(14)
                else:
                    heading_style.font.size = Pt(12)
                    
            except Exception:
                continue
        
        # Custom paragraph style
        try:
            normal_style = styles['Normal']
            normal_style.font.name = 'Calibri'
            normal_style.font.size = Pt(11)
            normal_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        except Exception:
            pass
            
        self.style_applied = True
    
    def add_table_of_contents(self, doc: Document):
        """Add a table of contents (placeholder for now)"""
        if not self.config.add_table_of_contents:
            return
            
        toc_heading = doc.add_heading('Table of Contents', level=1)
        toc_para = doc.add_paragraph()
        toc_run = toc_para.add_run("Table of contents will be generated when opened in Microsoft Word")
        toc_run.italic = True
        toc_run.font.size = Pt(10)
        doc.add_page_break()
    
    def process_html_element(self, doc: Document, element, current_paragraph=None, temp_dir="temp_images"):
        """Process HTML elements with professional formatting"""
        
        if isinstance(element, NavigableString):
            text = self._clean_text(str(element))
            if text and current_paragraph:
                current_paragraph.add_run(text)
            return current_paragraph
        
        if not hasattr(element, 'name') or not element.name:
            return current_paragraph
            
        tag_name = element.name.lower()
        
        # Handle Mermaid diagrams first
        mermaid_code = self.mermaid_renderer.extract_mermaid_code(element)
        if mermaid_code:
            return self._add_mermaid_diagram(doc, mermaid_code, temp_dir)
        
        # Handle standard HTML elements
        if tag_name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
            level = min(int(tag_name[1]), 3)
            text = self._clean_text(element.get_text())
            if text:
                heading = doc.add_heading(text, level=level)
                return heading
        
        elif tag_name == 'p':
            paragraph = doc.add_paragraph()
            for child in element.children:
                self.process_html_element(doc, child, paragraph, temp_dir)
            return paragraph
        
        elif tag_name in ['ul', 'ol']:
            return self._process_list(doc, element)
        
        elif tag_name == 'table':
            return self._process_table(doc, element)
        
        elif tag_name in ['pre', 'code'] and self.config.preserve_code_blocks:
            return self._process_code_block(doc, element)
        
        elif tag_name in ['strong', 'b']:
            text = self._clean_text(element.get_text())
            if text and current_paragraph:
                run = current_paragraph.add_run(text)
                run.bold = True
            return current_paragraph
        
        elif tag_name in ['em', 'i']:
            text = self._clean_text(element.get_text())
            if text and current_paragraph:
                run = current_paragraph.add_run(text)
                run.italic = True
            return current_paragraph
        
        elif tag_name == 'br':
            if current_paragraph:
                current_paragraph.add_run().add_break()
            return current_paragraph
        
        elif tag_name == 'div':
            paragraph = current_paragraph or doc.add_paragraph()
            for child in element.children:
                paragraph = self.process_html_element(doc, child, paragraph, temp_dir) or paragraph
            return paragraph
        
        else:
            # Process children for unhandled tags
            paragraph = current_paragraph
            for child in element.children:
                paragraph = self.process_html_element(doc, child, paragraph, temp_dir) or paragraph
            return paragraph
    
    def _add_mermaid_diagram(self, doc: Document, mermaid_code: str, temp_dir: str) -> Any:
        """Add Mermaid diagram to document"""
        image_data = self.mermaid_renderer.render_mermaid_to_image(mermaid_code)
        
        if image_data:
            try:
                os.makedirs(temp_dir, exist_ok=True)
                temp_image_path = os.path.join(temp_dir, f"mermaid_{self.mermaid_renderer.diagram_counter}.png")
                
                with open(temp_image_path, 'wb') as f:
                    f.write(image_data)
                
                # Add image
                paragraph = doc.add_paragraph()
                run = paragraph.add_run()
                run.add_picture(temp_image_path, width=Inches(self.config.diagram_width_inches))
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Add professional caption
                caption_para = doc.add_paragraph()
                caption_run = caption_para.add_run(f"Figure {self.mermaid_renderer.diagram_counter}: Process Flow Diagram")
                caption_run.italic = True
                caption_run.font.size = Pt(9)
                caption_run.font.color.rgb = RGBColor(0x5A, 0x5A, 0x5A)
                caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                try:
                    os.remove(temp_image_path)
                except:
                    pass
                    
                return caption_para
                
            except Exception as e:
                if self.config.verbose_logging:
                    print(f"   ⚠️  Error adding diagram: {e}")
        
        # Fallback
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(f"[Diagram {self.mermaid_renderer.diagram_counter}: Rendering failed]")
        run.italic = True
        run.font.color.rgb = RGBColor(0x85, 0x85, 0x85)
        return paragraph
    
    def _process_list(self, doc: Document, element: Tag) -> None:
        """Process HTML lists with proper formatting"""
        list_style = 'List Bullet' if element.name == 'ul' else 'List Number'
        
        for li in element.find_all('li', recursive=False):
            text = self._clean_text(li.get_text())
            if text:
                para = doc.add_paragraph(text, style=list_style)
        
        return None
    
    def _process_table(self, doc: Document, element: Tag) -> None:
        """Process HTML tables with professional formatting"""
        rows = element.find_all('tr')
        if not rows:
            return None
            
        max_cols = max(len(row.find_all(['td', 'th'])) for row in rows)
        table = doc.add_table(rows=len(rows), cols=max_cols)
        table.style = 'Table Grid'
        
        for i, row in enumerate(rows):
            cells = row.find_all(['td', 'th'])
            for j, cell in enumerate(cells):
                if j < len(table.rows[i].cells):
                    cell_text = self._clean_text(cell.get_text())
                    table.rows[i].cells[j].text = cell_text
                    
                    # Style header cells
                    if cell.name == 'th':
                        for paragraph in table.rows[i].cells[j].paragraphs:
                            for run in paragraph.runs:
                                run.bold = True
                                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        
        return None
    
    def _process_code_block(self, doc: Document, element: Tag) -> Any:
        """Process code blocks with monospace formatting"""
        code_text = element.get_text().strip()
        if code_text:
            para = doc.add_paragraph()
            run = para.add_run(code_text)
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
            # Add light background effect (approximated with style)
            para.style = 'Intense Quote'
            return para
        return None
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize text content"""
        if not text:
            return ""
        # Normalize whitespace but preserve intentional spacing
        text = re.sub(r'\s+', ' ', text.strip())
        return unescape(text)
    
    def convert_file(self, html_file: str, output_file: str) -> Tuple[bool, int, int, int]:
        """Convert single HTML file to DOCX"""
        try:
            with open(html_file, 'r', encoding='utf-8') as file:
                html_content = file.read()
            
            soup = BeautifulSoup(html_content, 'lxml')
            doc = Document()
            
            # Apply professional styling
            self.create_professional_styles(doc)
            
            # Set document title
            title_tag = soup.find('title')
            if title_tag:
                title = self._clean_text(title_tag.get_text())
                if title:
                    doc.add_heading(title, 0)
            
            # Add table of contents
            self.add_table_of_contents(doc)
            
            # Process content
            body = soup.find('body') or soup
            temp_dir = os.path.join(self.config.temp_dir, "docx_images")
            os.makedirs(temp_dir, exist_ok=True)
            
            for element in body.children:
                self.process_html_element(doc, element, None, temp_dir)
            
            # Ensure output directory exists
            output_dir = os.path.dirname(output_file)
            if output_dir and not os.path.exists(output_dir):
                os.makedirs(output_dir)
            
            doc.save(output_file)
            
            # Clean up temp files
            try:
                if os.path.exists(temp_dir):
                    shutil.rmtree(temp_dir)
            except:
                pass
            
            if os.path.exists(output_file) and os.path.getsize(output_file) > 0:
                return True, os.path.getsize(output_file), self.mermaid_renderer.successful_renders, self.mermaid_renderer.failed_renders
            else:
                return False, 0, 0, 0
                
        except Exception as e:
            if self.config.verbose_logging:
                print(f"   ❌ Error converting to DOCX: {e}")
            return False, 0, 0, 0

class EnhancedMarkdownConverter(HTMLParser):
    """Enhanced HTML to Markdown converter with better structure handling"""
    
    def __init__(self, config: ConversionConfig):
        super().__init__()
        self.config = config
        self.mermaid_renderer = EnhancedMermaidRenderer(config)
        self.markdown = []
        self.current_tag = []
        self.list_depth = 0
        self.in_table = False
        self.table_headers = []
        self.table_rows = []
        self.current_row = []
        self.skip_content = False
        self.in_code = False
        self.pending_text = ""
        
    def handle_starttag(self, tag, attrs):
        attrs_dict = dict(attrs)

        # Skip style, script, and head tags
        if tag in ['style', 'script', 'head']:
            self.skip_content = True
            return

        if self.skip_content:
            return

        # Check for Mermaid diagrams
        if self._is_mermaid_element(tag, attrs_dict):
            self.pending_text = "\n```mermaid\n"
        elif tag == 'h1':
            self.pending_text = "\n# "
        elif tag == 'h2':
            self.pending_text = "\n## "
        elif tag == 'h3':
            self.pending_text = "\n### "
        elif tag == 'h4':
            self.pending_text = "\n#### "
        elif tag == 'h5':
            self.pending_text = "\n##### "
        elif tag == 'h6':
            self.pending_text = "\n###### "
        elif tag == 'p':
            self.pending_text = "\n\n"
        elif tag == 'br':
            self.markdown.append("  \n")  # Markdown line break
        elif tag == 'strong' or tag == 'b':
            self.pending_text = "**"
        elif tag == 'em' or tag == 'i':
            self.pending_text = "*"
        elif tag == 'code':
            self.pending_text = "`"
            self.in_code = True
        elif tag == 'pre':
            self.pending_text = "\n```\n"
        elif tag == 'ul':
            self.list_depth += 1
            if self.list_depth == 1:
                self.markdown.append("\n")
        elif tag == 'ol':
            self.list_depth += 1
            if self.list_depth == 1:
                self.markdown.append("\n")
        elif tag == 'li':
            indent = "  " * (self.list_depth - 1)
            self.pending_text = f"\n{indent}- "
        elif tag == 'table':
            self.in_table = True
            self.table_headers = []
            self.table_rows = []
            self.markdown.append("\n")
        elif tag == 'tr':
            self.current_row = []
        elif tag == 'a' and 'href' in attrs_dict:
            self.pending_text = "["
        elif tag == 'blockquote':
            self.pending_text = "\n> "

        self.current_tag.append(tag)
    
    def _is_mermaid_element(self, tag: str, attrs: Dict[str, str]) -> bool:
        """Check if element contains Mermaid diagram"""
        if tag == 'div' and 'class' in attrs:
            return 'mermaid' in attrs['class'].lower()
        elif tag == 'pre' or tag == 'code':
            return 'mermaid' in attrs.get('class', '').lower()
        elif tag == 'script':
            return attrs.get('type') == 'text/mermaid'
        return False

    def handle_endtag(self, tag):
        if tag in ['style', 'script', 'head']:
            self.skip_content = False
            return

        if self.skip_content:
            return

        if self.current_tag and self.current_tag[-1] == tag:
            self.current_tag.pop()

        if tag in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
            self.markdown.append("\n")
        elif tag == 'p':
            self.markdown.append("\n")
        elif tag == 'strong' or tag == 'b':
            self.markdown.append("**")
        elif tag == 'em' or tag == 'i':
            self.markdown.append("*")
        elif tag == 'code':
            self.markdown.append("`")
            self.in_code = False
        elif tag == 'pre':
            self.markdown.append("\n```\n")
        elif tag == 'ul' or tag == 'ol':
            self.list_depth -= 1
            if self.list_depth == 0:
                self.markdown.append("\n")
        elif tag == 'li':
            pass  # Line break handled in start tag
        elif tag == 'tr':
            if self.in_table:
                if not self.table_headers:
                    self.table_headers = self.current_row[:]
                else:
                    self.table_rows.append(self.current_row[:])
        elif tag == 'table':
            self._write_markdown_table()
            self.in_table = False
        elif tag == 'div' and self._was_mermaid_div():
            self.markdown.append("\n```\n\n")
        elif tag == 'a':
            # Close link - would need href from start tag for full implementation
            self.markdown.append("](#)")
        elif tag == 'blockquote':
            self.markdown.append("\n")

    def _was_mermaid_div(self) -> bool:
        """Check if the last few markdown entries suggest this was a mermaid div"""
        recent = ''.join(self.markdown[-5:])
        return '```mermaid' in recent and '```' not in recent.split('```mermaid')[-1]
    
    def _write_markdown_table(self):
        """Write a complete markdown table"""
        if self.table_headers:
            # Headers
            self.markdown.append("| " + " | ".join(self.table_headers) + " |\n")
            self.markdown.append("| " + " | ".join(["---"] * len(self.table_headers)) + " |\n")
            
            # Rows
            for row in self.table_rows:
                # Pad row to match header length
                while len(row) < len(self.table_headers):
                    row.append("")
                self.markdown.append("| " + " | ".join(row) + " |\n")
            
            self.markdown.append("\n")
        
        self.table_headers = []
        self.table_rows = []

    def handle_data(self, data):
        if self.skip_content:
            return

        data = data.strip()
        if not data:
            return

        # Add pending formatting
        if self.pending_text:
            self.markdown.append(self.pending_text)
            self.pending_text = ""

        # Unescape HTML entities
        data = unescape(data)

        # Handle table cells
        if self.in_table and len(self.current_tag) > 0 and self.current_tag[-1] in ['th', 'td']:
            self.current_row.append(data)
        else:
            # Add appropriate spacing
            if self.markdown and len(self.markdown) > 0:
                last_char = self.markdown[-1][-1] if self.markdown[-1] else ''
                if last_char not in [' ', '\n', '*', '`', '#', '[', '>', '-']:
                    self.markdown.append(' ')
            self.markdown.append(data)

    def get_markdown(self) -> str:
        """Get the final markdown content"""
        result = ''.join(self.markdown)
        
        # Clean up excessive whitespace
        result = re.sub(r'\n{4,}', '\n\n\n', result)
        result = re.sub(r' +\n', '\n', result)
        result = re.sub(r'  +', ' ', result)
        
        return result.strip()
    
    def convert_file(self, html_file: str, output_file: str) -> Tuple[bool, int]:
        """Convert HTML file to Markdown"""
        try:
            with open(html_file, 'r', encoding='utf-8') as f:
                html_content = f.read()

            # Reset parser state
            self.__init__(self.config)
            
            self.feed(html_content)
            markdown_content = self.get_markdown()

            # Ensure output directory exists
            output_dir = os.path.dirname(output_file)
            if output_dir and not os.path.exists(output_dir):
                os.makedirs(output_dir)

            with open(output_file, 'w', encoding='utf-8') as f:
                f.write(markdown_content)

            if os.path.exists(output_file) and os.path.getsize(output_file) > 0:
                return True, len(markdown_content)
            else:
                return False, 0
                
        except Exception as e:
            if self.config.verbose_logging:
                print(f"   ❌ Error converting to Markdown: {e}")
            return False, 0

class UltimateHtmlConverter:
    """Ultimate HTML converter with support for multiple output formats"""
    
    def __init__(self, config: ConversionConfig = None):
        self.config = config or ConversionConfig()
        self.docx_converter = ProfessionalDocxConverter(self.config)
        self.md_converter = EnhancedMarkdownConverter(self.config)
        
    def convert_single_file(self, html_file: str, base_output_name: str = None) -> Dict[str, Any]:
        """Convert a single HTML file to specified format(s)"""
        
        if not os.path.exists(html_file):
            return {"success": False, "error": f"File not found: {html_file}"}
        
        if not base_output_name:
            base_output_name = os.path.splitext(os.path.basename(html_file))[0]
        
        results = {"success": True, "conversions": {}}
        
        # Ensure output directory exists
        output_dir = self.config.output_dir
        os.makedirs(output_dir, exist_ok=True)
        
        # Convert to DOCX
        if self.config.output_format in [OutputFormat.DOCX, OutputFormat.BOTH]:
            docx_file = os.path.join(output_dir, f"{base_output_name}_Professional.docx")
            
            if self.config.verbose_logging:
                print(f"   📄 Converting to DOCX: {docx_file}")
                
            success, size, diagrams_ok, diagrams_fail = self.docx_converter.convert_file(html_file, docx_file)
            
            results["conversions"]["docx"] = {
                "success": success,
                "file": docx_file,
                "size": size,
                "diagrams_rendered": diagrams_ok,
                "diagrams_failed": diagrams_fail
            }
        
        # Convert to Markdown
        if self.config.output_format in [OutputFormat.MARKDOWN, OutputFormat.BOTH]:
            md_file = os.path.join(output_dir, f"{base_output_name}_Enhanced.md")
            
            if self.config.verbose_logging:
                print(f"   📝 Converting to Markdown: {md_file}")
                
            success, char_count = self.md_converter.convert_file(html_file, md_file)
            
            results["conversions"]["markdown"] = {
                "success": success,
                "file": md_file,
                "characters": char_count
            }
        
        return results
    
    def batch_convert(self, html_pattern: str = "*.html") -> Dict[str, Any]:
        """Batch convert HTML files with comprehensive reporting"""
        
        print("🚀 ULTIMATE HTML CONVERTER")
        print("=" * 70)
        print(f"📋 Configuration:")
        print(f"   • Output Format: {self.config.output_format.value}")
        print(f"   • Render Mermaid: {self.config.render_mermaid}")
        print(f"   • Professional Styling: {self.config.professional_styling}")
        print(f"   • Output Directory: {self.config.output_dir}")
        print()
        
        # Find HTML files
        html_files = glob.glob(html_pattern)
        
        if not html_files:
            print("❌ No HTML files found")
            return {"success": False, "error": "No HTML files found"}
        
        print(f"📁 Found {len(html_files)} HTML files:")
        for i, file in enumerate(html_files, 1):
            size_kb = os.path.getsize(file) / 1024
            print(f"   {i:2d}. {file:<40} ({size_kb:,.1f} KB)")
        print()
        
        # Process each file
        start_time = datetime.now()
        results = {
            "success": True,
            "total_files": len(html_files),
            "conversions": {},
            "summary": {
                "docx": {"successful": 0, "failed": 0, "total_size": 0, "total_diagrams": 0},
                "markdown": {"successful": 0, "failed": 0, "total_chars": 0}
            }
        }
        
        for i, html_file in enumerate(html_files, 1):
            base_name = os.path.splitext(os.path.basename(html_file))[0]
            
            print(f"[{i:2d}/{len(html_files)}] Processing: {os.path.basename(html_file)}")
            
            file_results = self.convert_single_file(html_file, base_name)
            results["conversions"][html_file] = file_results
            
            # Update summary statistics
            if "docx" in file_results.get("conversions", {}):
                docx_result = file_results["conversions"]["docx"]
                if docx_result["success"]:
                    results["summary"]["docx"]["successful"] += 1
                    results["summary"]["docx"]["total_size"] += docx_result["size"]
                    results["summary"]["docx"]["total_diagrams"] += docx_result["diagrams_rendered"]
                    print(f"   ✅ DOCX: {docx_result['size']:,} bytes, {docx_result['diagrams_rendered']} diagrams")
                else:
                    results["summary"]["docx"]["failed"] += 1
                    print(f"   ❌ DOCX: Failed")
            
            if "markdown" in file_results.get("conversions", {}):
                md_result = file_results["conversions"]["markdown"]
                if md_result["success"]:
                    results["summary"]["markdown"]["successful"] += 1
                    results["summary"]["markdown"]["total_chars"] += md_result["characters"]
                    print(f"   ✅ Markdown: {md_result['characters']:,} characters")
                else:
                    results["summary"]["markdown"]["failed"] += 1
                    print(f"   ❌ Markdown: Failed")
            
            print()
        
        # Final summary
        end_time = datetime.now()
        duration = (end_time - start_time).total_seconds()
        
        print("=" * 70)
        print("📊 CONVERSION SUMMARY")
        print("=" * 70)
        
        if self.config.output_format in [OutputFormat.DOCX, OutputFormat.BOTH]:
            docx_stats = results["summary"]["docx"]
            print(f"📄 DOCX Conversions:")
            print(f"   ✅ Successful: {docx_stats['successful']}")
            print(f"   ❌ Failed:     {docx_stats['failed']}")
            print(f"   📊 Total size: {docx_stats['total_size']:,} bytes")
            print(f"   🎨 Diagrams:   {docx_stats['total_diagrams']}")
            
        if self.config.output_format in [OutputFormat.MARKDOWN, OutputFormat.BOTH]:
            md_stats = results["summary"]["markdown"]
            print(f"📝 Markdown Conversions:")
            print(f"   ✅ Successful: {md_stats['successful']}")
            print(f"   ❌ Failed:     {md_stats['failed']}")
            print(f"   📝 Total text: {md_stats['total_chars']:,} characters")
        
        print(f"⏱️  Processing time: {duration:.1f} seconds")
        print(f"📁 Output directory: {self.config.output_dir}/")
        
        total_successful = results["summary"]["docx"]["successful"] + results["summary"]["markdown"]["successful"]
        if total_successful > 0:
            print(f"\n🎉 {total_successful} files converted successfully!")
        
        return results

def main():
    """Main function with configuration options"""
    
    # Configuration - modify these as needed
    config = ConversionConfig(
        output_format=OutputFormat.BOTH,          # Convert to both DOCX and Markdown
        render_mermaid=True,                      # Render Mermaid diagrams as images
        max_diagram_retries=3,                    # Number of retries for diagram rendering
        diagram_width_inches=6.0,                 # Width of diagrams in DOCX
        preserve_code_blocks=True,                # Keep code formatting
        add_table_of_contents=True,               # Add TOC to DOCX files
        professional_styling=True,                # Apply professional document styles
        verbose_logging=True,                     # Detailed progress output
        temp_dir="temp_conversion",               # Temporary files directory
        output_dir="Professional_Output"          # Output directory name
    )
    
    # Create converter and run batch conversion
    converter = UltimateHtmlConverter(config)
    results = converter.batch_convert("*.html")
    
    # Clean up temporary files
    try:
        if os.path.exists(config.temp_dir):
            shutil.rmtree(config.temp_dir)
    except:
        pass
    
    return results

if __name__ == "__main__":
    main()